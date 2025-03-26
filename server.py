from flask import Flask, request, jsonify, Response, stream_with_context, send_from_directory
from flask_cors import CORS
from helper_function import create_anthropic_client, create_gemini_client, GeminiStreamingResponse
import anthropic
import json
import os
import re
import time
import traceback
# Updated import to be compatible with different versions of the Anthropic library
try:
    from anthropic.types import TextBlock, MessageParam
except ImportError:
    # Fallback for newer versions of the Anthropic library
    # where these classes might have different names or locations
    TextBlock = dict
    MessageParam = dict
import uuid
import PyPDF2
import docx
import io
import base64
import random 
import socket
import argparse
import logging
from datetime import datetime

# Import version
try:
    from . import __version__
except ImportError:
    __version__ = "0.4.5"  # Fallback version

# Import Google Generative AI package
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("Google Generative AI package not available. Some features may be limited.")

# Initialize Flask app
app = Flask(__name__, static_folder='static')
CORS(app)  # Enable CORS for all routes

# Set higher request timeout limits for Flask server
app.config['TIMEOUT'] = 1800  # 30 minutes timeout
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max content size

# Simple in-memory session cache (for production, consider Redis)
session_cache = {}
SESSION_CACHE_EXPIRY = 3600  # 1 hour cache expiry

# Claude 3.7 has a total context window of 200,000 tokens (input + output combined)
# We'll use this constant when estimating token usage
TOTAL_CONTEXT_WINDOW = 200000

# Maximum allowed tokens for Claude API input (leaving room for output)
# This should be dynamically adjusted based on requested max_tokens
MAX_INPUT_TOKENS = 195000  # Setting a bit lower than the actual limit to account for system prompt and overhead

# Default settings
DEFAULT_MAX_TOKENS = 128000
DEFAULT_THINKING_BUDGET = 32000  # Kept for compatibility but thinking tokens are included in output tokens
STREAM_CHUNK_SIZE = 2  # Send keepalive every 2 chunks (reduced from 5)
MAX_SEGMENT_SIZE = 16384  # 16KB per segment (reduced from 32KB)

# Define beta parameter for 128K output
OUTPUT_128K_BETA = "output-128k-2025-02-19"

# Set higher request timeout and stream chunk sizes
MAX_TOKENS = 4096
STREAM_CHUNK_SIZE = 2  # Number of chunks to process before sending a keepalive
MAX_SEGMENT_SIZE = 16384  # 16KB chunks for content segments
CHECKPOINT_INTERVAL = 2 * 60  # 2 minutes between checkpoints (reduced from 5)

# Retry settings
MAX_RETRIES = 10  # Increase from 8 to 10
MIN_BACKOFF_DELAY = 1  # Start with 1 second delay (reduced from 2)
MAX_BACKOFF_DELAY = 45  # Max 45 seconds delay (reduced from 60)
BACKOFF_FACTOR = 1.3  # Use 1.3 instead of 1.5 for more gradual increase

# Gemini-specific settings
GEMINI_MODEL = "gemini-2.5-pro-exp-03-25"
GEMINI_MAX_OUTPUT_TOKENS = 65536
GEMINI_TEMPERATURE = 1.0
GEMINI_TOP_P = 0.95
GEMINI_TOP_K = 64

# Same system instruction for both APIs
SYSTEM_INSTRUCTION = """I will provide you with a file or a content, analyze its content, and transform it into a visually appealing and well-structured webpage.### Content Requirements* Maintain the core information from the original file while presenting it in a clearer and more visually engaging format.⠀Design Style* Follow a modern and minimalistic design inspired by Linear App.* Use a clear visual hierarchy to emphasize important content.* Adopt a professional and harmonious color scheme that is easy on the eyes for extended reading.⠀Technical Specifications* Use HTML5, TailwindCSS 3.0+ (via CDN), and necessary JavaScript.* Implement a fully functional dark/light mode toggle, defaulting to the system setting.* Ensure clean, well-structured code with appropriate comments for easy understanding and maintenance.⠀Responsive Design* The page must be fully responsive, adapting seamlessly to mobile, tablet, and desktop screens.* Optimize layout and typography for different screen sizes.* Ensure a smooth and intuitive touch experience on mobile devices.⠀Icons & Visual Elements* Use professional icon libraries like Font Awesome or Material Icons (via CDN).* Integrate illustrations or charts that best represent the content.* Avoid using emojis as primary icons.* Check if any icons cannot be loaded.⠀User Interaction & ExperienceEnhance the user experience with subtle micro-interactions:* Buttons should have slight enlargement and color transitions on hover.* Cards should feature soft shadows and border effects on hover.* Implement smooth scrolling effects throughout the page.* Content blocks should have an elegant fade-in animation on load.⠀Performance Optimization* Ensure fast page loading by avoiding large, unnecessary resources.* Use modern image formats (WebP) with proper compression.* Implement lazy loading for content-heavy pages.⠀Output Requirements* Deliver a fully functional standalone HTML file, including all necessary CSS and JavaScript.* Ensure the code meets W3C standards with no errors or warnings.* Maintain consistent design and functionality across different browsers.* Your output is only one HTML file, do not present any other notes on the HTML. Also, try your best to visualize the whole content.⠀Create the most effective and visually appealing webpage based on the uploaded file's content type (document, data, images, etc.)."""

@app.route('/')
def serve_index():
    """Serve the main index.html file with version information in headers."""
    response = send_from_directory('static', 'index.html')
    response.headers['X-App-Version'] = __version__
    return response

@app.route('/<path:path>')
def static_files(path):
    return send_from_directory('static', path)

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    try:
        # Read the file content
        file_content = file.read()
        # Convert the file content to base64
        base64_file_content = base64.b64encode(file_content).decode('utf-8')
        # Create a response object
        response = {
            "file_name": file.filename,
            "file_content": base64_file_content
        }
        return jsonify(response)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/process', methods=['POST'])
def process():
    data = request.get_json()
    if not data or 'file_name' not in data or 'file_content' not in data:
        return jsonify({"error": "Missing file_name or file_content"}), 400

    file_name = data['file_name']
    file_content = data['file_content']
    
    # Get additional parameters from request or use defaults
    api_key = data.get('api_key', '')
    format_prompt = data.get('format_prompt', '')
    model = data.get('model', 'claude-3-7-sonnet-20250219')
    max_tokens = int(data.get('max_tokens', DEFAULT_MAX_TOKENS))
    temperature = float(data.get('temperature', 1.0))
    thinking_budget = int(data.get('thinking_budget', DEFAULT_THINKING_BUDGET))

    try:
        # Convert base64 string back to bytes
        file_content_bytes = base64.b64decode(file_content)
        # Create a temporary file to save the uploaded file
        temp_file_path = f"/tmp/{file_name}"
        with open(temp_file_path, 'wb') as f:
            f.write(file_content_bytes)
            
        # Determine file type and read content
        file_ext = file_name.split('.')[-1].lower() if '.' in file_name else 'txt'
        
        # Process the file based on its type
        file_text_content = ""
        
        if file_ext == 'pdf':
            # Process PDF file
            try:
                reader = PyPDF2.PdfReader(temp_file_path)
                for page_num in range(len(reader.pages)):
                    file_text_content += reader.pages[page_num].extract_text() + "\n"
            except Exception as e:
                return jsonify({"error": f"Error processing PDF: {str(e)}"}), 500
                
        elif file_ext in ['docx', 'doc']:
            # Process Word document
            try:
                doc = docx.Document(temp_file_path)
                file_text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as e:
                return jsonify({"error": f"Error processing Word document: {str(e)}"}), 500
                
        else:
            # Process text-based file
            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                file_text_content = f.read()
                
        # Now process the file content with Claude
        if not api_key:
            return jsonify({"error": "API key is required"}), 400
            
        try:
            # Use our helper function to create a compatible client
            client = create_anthropic_client(api_key)
            
            # Prepare user message with content and additional prompt
            user_content = file_text_content
            if format_prompt:
                user_content = f"{user_content}\n\n{format_prompt}"
            
            # Create parameters for the API call
            params = {
                "model": model,
                "max_tokens": max_tokens,
                "temperature": temperature,
                "system": "I will provide you with a file or a content, analyze its content, and transform it into a visually appealing and well-structured webpage.### Content Requirements* Maintain the core information from the original file while presenting it in a clearer and more visually engaging format.⠀Design Style* Follow a modern and minimalistic design inspired by Linear App.* Use a clear visual hierarchy to emphasize important content.* Adopt a professional and harmonious color scheme that is easy on the eyes for extended reading.⠀Technical Specifications* Use HTML5, TailwindCSS 3.0+ (via CDN), and necessary JavaScript.* Implement a fully functional dark/light mode toggle, defaulting to the system setting.* Ensure clean, well-structured code with appropriate comments for easy understanding and maintenance.⠀Responsive Design* The page must be fully responsive, adapting seamlessly to mobile, tablet, and desktop screens.* Optimize layout and typography for different screen sizes.* Ensure a smooth and intuitive touch experience on mobile devices.⠀Icons & Visual Elements* Use professional icon libraries like Font Awesome or Material Icons (via CDN).* Integrate illustrations or charts that best represent the content.* Avoid using emojis as primary icons.* Check if any icons cannot be loaded.⠀User Interaction & ExperienceEnhance the user experience with subtle micro-interactions:* Buttons should have slight enlargement and color transitions on hover.* Cards should feature soft shadows and border effects on hover.* Implement smooth scrolling effects throughout the page.* Content blocks should have an elegant fade-in animation on load.⠀Performance Optimization* Ensure fast page loading by avoiding large, unnecessary resources.* Use modern image formats (WebP) with proper compression.* Implement lazy loading for content-heavy pages.⠀Output Requirements* Deliver a fully functional standalone HTML file, including all necessary CSS and JavaScript.* Ensure the code meets W3C standards with no errors or warnings.* Maintain consistent design and functionality across different browsers.* Your output is only one HTML file, do not present any other notes on the HTML. Also, try your best to visualize the whole content.⠀Create the most effective and visually appealing webpage based on the uploaded file's content type (document, data, images, etc.).",
                "messages": [{"role": "user", "content": user_content}],
            }
            
            # Add thinking parameter if thinking_budget > 0
            if thinking_budget > 0:
                params["thinking"] = {"type": "enabled", "budget_tokens": thinking_budget}
                
            # Add beta parameter if needed for large outputs
            if max_tokens > 4096:
                params["betas"] = [OUTPUT_128K_BETA]
            
            # Create the message
            response = client.messages.create(**params)
            
            # Extract the HTML from the response
            html_content = ""
            
            # Handle different response formats
            if hasattr(response, 'content') and response.content:
                # Handle standard Anthropic client response
                if isinstance(response.content, list) and len(response.content) > 0:
                    if hasattr(response.content[0], 'text'):
                        # Standard client format
                        html_content = response.content[0].text
                    elif isinstance(response.content[0], dict) and 'text' in response.content[0]:
                        # Dictionary format
                        html_content = response.content[0]['text']
                elif isinstance(response.content, dict) and 'text' in response.content:
                    # Simple dictionary format
                    html_content = response.content['text']
                elif isinstance(response.content, str):
                    # Direct string format
                    html_content = response.content
            
            # If still empty, try alternate paths
            if not html_content and hasattr(response, 'completion'):
                html_content = response.completion
                
            # As a last resort, convert the entire response to string
            if not html_content:
                print("Warning: Unable to extract HTML content using standard methods. Using fallback extraction.")
                try:
                    # Try to extract from the raw response
                    if isinstance(response, dict) and 'content' in response:
                        if isinstance(response['content'], list) and len(response['content']) > 0:
                            first_item = response['content'][0]
                            if isinstance(first_item, dict) and 'text' in first_item:
                                html_content = first_item['text']
                    
                    # If still empty, convert the whole response to string
                    if not html_content:
                        html_content = str(response)
                except Exception as e:
                    print(f"Fallback extraction failed: {str(e)}")
                    html_content = "Error: Unable to extract HTML content from response."
                
            # Get usage stats
            input_tokens = 0
            output_tokens = 0
            
            if hasattr(response, 'usage'):
                if hasattr(response.usage, 'input_tokens'):
                    input_tokens = response.usage.input_tokens
                if hasattr(response.usage, 'output_tokens'):
                    output_tokens = response.usage.output_tokens
            elif isinstance(response, dict) and 'usage' in response:
                usage = response['usage']
                if isinstance(usage, dict):
                    input_tokens = usage.get('input_tokens', 0)
                    output_tokens = usage.get('output_tokens', 0)
            
            # Calculate total cost based on token usage
            # Based on Anthropic documentation: Claude 3.7 Sonnet costs $3/MTok for input and $15/MTok for output
            # Thinking tokens are already included in output tokens
            total_cost = (input_tokens / 1000000) * 3.0 + (output_tokens / 1000000) * 15.0
                
            # Return the response
            return jsonify({
                'html': html_content,
                'model': model,
                'usage': {
                    'input_tokens': input_tokens,
                    'output_tokens': output_tokens,
                    'total_cost': total_cost
                }
            })
            
        except Exception as e:
            error_message = str(e)
            print(f"Error processing file content with Claude: {error_message}")
            return jsonify({'error': f'Claude API error: {error_message}'}), 500
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/validate-key', methods=['POST'])
def validate_key():
    print("\n==== API VALIDATE KEY REQUEST RECEIVED ====")
    data = request.get_json()
    
    if not data or 'api_key' not in data:
        return jsonify({"valid": False, "message": "API key is required"}), 400
    
    api_key = data['api_key']
    api_type = data.get('api_type', 'anthropic')  # Default to Anthropic for backward compatibility
    print(f"Request JSON: {data}")
    
    # Basic format validation
    if not api_key or not api_key.strip():
        return jsonify({"valid": False, "message": "API key cannot be empty"}), 400
    
    # Different validation based on API type
    if api_type == 'anthropic':
        # Check if the API key has a valid format (starts with sk-ant)
        if not api_key.startswith('sk-ant'):
            return jsonify({
                "valid": False, 
                "message": "Anthropic API key format is invalid. It should start with 'sk-ant'"
            }), 400
        
        print(f"Anthropic API key format is valid: {api_key[:10]}...")
        
        try:
            # Try to create a client to validate the key
            client = create_anthropic_client(api_key)
            
            # For security, we don't actually make an API call here
            # Just successfully creating the client is enough validation
            
            return jsonify({
                "valid": True,
                "message": "Anthropic API key is valid"
            })
        except Exception as e:
            error_message = str(e)
            print(f"Anthropic API key validation error: {error_message}")
            
            # Check for specific error messages
            if "proxies" in error_message.lower():
                print(f"Proxies error detected: {error_message}")
                # This is a client configuration issue, not an invalid key
                return jsonify({
                    "valid": True,  # Consider the key valid despite the client error
                    "message": "Anthropic API key format is valid, but there was a client configuration issue"
                })
            
            return jsonify({
                "valid": False,
                "message": "Anthropic API client configuration error. Please try using a newer Anthropic API key format."
            }), 400
    
    elif api_type == 'gemini':
        # For Gemini, we'll check if the package is available first
        if not GEMINI_AVAILABLE:
            return jsonify({
                "valid": False,
                "message": "Google Generative AI package is not installed on the server."
            }), 500
        
        # Gemini API keys don't have a specific format to validate upfront
        # We'll try to create a client and check if it works
        try:
            # Try to create a client to validate the key
            client = create_gemini_client(api_key)
            
            # For security, we don't actually make an API call here
            # Just successfully creating the client is enough validation
            
            return jsonify({
                "valid": True,
                "message": "Google Gemini API key is valid"
            })
        except Exception as e:
            error_message = str(e)
            print(f"Google Gemini API key validation error: {error_message}")
            
            return jsonify({
                "valid": False,
                "message": f"Google Gemini API error: {error_message}"
            }), 400
    
    else:
        return jsonify({
            "valid": False,
            "message": f"Unsupported API type: {api_type}"
        }), 400

@app.route('/api/process', methods=['POST'])
def process_file():
    # Get the data from the request
    print("\n==== API PROCESS REQUEST RECEIVED ====")
    data = request.get_json()
    
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    # Extract the API key and content
    api_key = data.get('api_key')
    content = data.get('content')
    format_prompt = data.get('format_prompt', '')
    model = data.get('model', 'claude-3-7-sonnet-20250219')
    max_tokens = int(data.get('max_tokens', 128000))
    temperature = float(data.get('temperature', 1.0))
    thinking_budget = int(data.get('thinking_budget', 32000))
    
    print(f"Processing request with model={model}, max_tokens={max_tokens}, content_length={len(content) if content else 0}")
    
    # Check if we have the required data
    if not api_key or not content:
        return jsonify({'error': 'API key and content are required'}), 400
    
    try:
        # Use our helper function to create a compatible client
        client = create_anthropic_client(api_key)
        
        # Prepare user message with content and additional prompt
        user_content = content
        if format_prompt:
            user_content = f"{user_content}\n\n{format_prompt}"
        
        print("Creating message with thinking parameter...")
        
        # Create parameters for the API call
        params = {
            "model": model,
            "max_tokens": max_tokens,
            "temperature": temperature,
            "system": "I will provide you with a file or a content, analyze its content, and transform it into a visually appealing and well-structured webpage.### Content Requirements* Maintain the core information from the original file while presenting it in a clearer and more visually engaging format.⠀Design Style* Follow a modern and minimalistic design inspired by Linear App.* Use a clear visual hierarchy to emphasize important content.* Adopt a professional and harmonious color scheme that is easy on the eyes for extended reading.⠀Technical Specifications* Use HTML5, TailwindCSS 3.0+ (via CDN), and necessary JavaScript.* Implement a fully functional dark/light mode toggle, defaulting to the system setting.* Ensure clean, well-structured code with appropriate comments for easy understanding and maintenance.⠀Responsive Design* The page must be fully responsive, adapting seamlessly to mobile, tablet, and desktop screens.* Optimize layout and typography for different screen sizes.* Ensure a smooth and intuitive touch experience on mobile devices.⠀Icons & Visual Elements* Use professional icon libraries like Font Awesome or Material Icons (via CDN).* Integrate illustrations or charts that best represent the content.* Avoid using emojis as primary icons.* Check if any icons cannot be loaded.⠀User Interaction & ExperienceEnhance the user experience with subtle micro-interactions:* Buttons should have slight enlargement and color transitions on hover.* Cards should feature soft shadows and border effects on hover.* Implement smooth scrolling effects throughout the page.* Content blocks should have an elegant fade-in animation on load.⠀Performance Optimization* Ensure fast page loading by avoiding large, unnecessary resources.* Use modern image formats (WebP) with proper compression.* Implement lazy loading for content-heavy pages.⠀Output Requirements* Deliver a fully functional standalone HTML file, including all necessary CSS and JavaScript.* Ensure the code meets W3C standards with no errors or warnings.* Maintain consistent design and functionality across different browsers.⠀Create the most effective and visually appealing webpage based on the uploaded file's content type (document, data, images, etc.Your output is only one HTML file, do not present any other notes on the HTML.",
            "messages": [{"role": "user", "content": user_content}],
        }
        
        # Add thinking parameter if thinking_budget > 0
        if thinking_budget > 0:
            params["thinking"] = {"type": "enabled", "budget_tokens": thinking_budget}
            
        # Add beta parameter if needed
        if max_tokens > 4096:
            params["betas"] = [OUTPUT_128K_BETA]
        
        # Create the message
        response = client.messages.create(**params)
        
        # Extract the HTML from the response
        html_content = ""
        
        # Handle different response formats
        if hasattr(response, 'content') and response.content:
            # Handle standard Anthropic client response
            if isinstance(response.content, list) and len(response.content) > 0:
                if hasattr(response.content[0], 'text'):
                    # Standard client format
                    html_content = response.content[0].text
                elif isinstance(response.content[0], dict) and 'text' in response.content[0]:
                    # Dictionary format
                    html_content = response.content[0]['text']
            elif isinstance(response.content, dict) and 'text' in response.content:
                # Simple dictionary format
                html_content = response.content['text']
            elif isinstance(response.content, str):
                # Direct string format
                html_content = response.content
        
        # If still empty, try alternate paths
        if not html_content and hasattr(response, 'completion'):
            html_content = response.completion
            
        # As a last resort, convert the entire response to string
        if not html_content:
            print("Warning: Unable to extract HTML content using standard methods. Using fallback extraction.")
            try:
                # Try to extract from the raw response
                if isinstance(response, dict) and 'content' in response:
                    if isinstance(response['content'], list) and len(response['content']) > 0:
                        first_item = response['content'][0]
                        if isinstance(first_item, dict) and 'text' in first_item:
                            html_content = first_item['text']
                
                # If still empty, convert the whole response to string
                if not html_content:
                    html_content = str(response)
            except Exception as e:
                print(f"Fallback extraction failed: {str(e)}")
                html_content = "Error: Unable to extract HTML content from response."
            
        # Get usage stats
        input_tokens = 0
        output_tokens = 0
        
        if hasattr(response, 'usage'):
            if hasattr(response.usage, 'input_tokens'):
                input_tokens = response.usage.input_tokens
            if hasattr(response.usage, 'output_tokens'):
                output_tokens = response.usage.output_tokens
        elif isinstance(response, dict) and 'usage' in response:
            usage = response['usage']
            if isinstance(usage, dict):
                input_tokens = usage.get('input_tokens', 0)
                output_tokens = usage.get('output_tokens', 0)
                
        # Log response structure for debugging
        print(f"Response type: {type(response)}")
        if html_content:
            print(f"Successfully extracted HTML content. Length: {len(html_content)}")
                
        # Return the response
        return jsonify({
            'html': html_content,
            'model': model,
            'usage': {
                'input_tokens': input_tokens,
                'output_tokens': output_tokens,
                'total_cost': (input_tokens / 1000000) * 3.0 + (output_tokens / 1000000) * 15.0
            }
        })
    
    except Exception as e:
        error_message = str(e)
        print(f"Error in /api/process: {error_message}")
        return jsonify({'error': f'Server error: {error_message}'}), 500

@app.route('/api/analyze-tokens', methods=['POST'])
def analyze_tokens():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400
        
        print(f"Analyzing tokens for data of size: {len(str(data))}")
        
        content = data.get('content', '')
        if not content:
            content = data.get('source', '')
            
        file_type = data.get('file_type', 'txt')
        
        # Handle PDF files (which are sent as base64)
        if file_type == 'pdf':
            # Recognize base64 data (could start with data:application/pdf;base64, or just be raw base64)
            if ';base64,' in content:
                # Extract the base64 part
                content = content.split(';base64,')[1]
            
            try:
                # Decode base64
                pdf_data = base64.b64decode(content)
                
                # Create a file-like object
                pdf_file = io.BytesIO(pdf_data)
                
                # Extract text from the PDF
                reader = PyPDF2.PdfReader(pdf_file)
                text_content = ""
                
                # Extract text from each page
                for page_num in range(len(reader.pages)):
                    text_content += reader.pages[page_num].extract_text() + "\n"
                
                # Use the extracted text for token analysis
                content = text_content
                
            except Exception as e:
                return jsonify({"error": f"Error processing PDF: {str(e)}"}), 400
        
        # Handle DOCX files
        elif file_type in ['docx', 'doc']:
            try:
                # Decode base64
                docx_data = base64.b64decode(content)
                
                # Create a file-like object
                docx_file = io.BytesIO(docx_data)
                
                # Extract text from the DOCX
                doc = docx.Document(docx_file)
                text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                # Use the extracted text for token analysis
                content = text_content
                
            except Exception as e:
                return jsonify({"error": f"Error processing DOCX: {str(e)}"}), 400
        
        # If no content after processing, return an error
        if not content:
            return jsonify({"error": "No content to analyze"}), 400
        
        # Define the system prompt to include in token estimation - Use the same detailed prompt as in process_file
        system_prompt = "I will provide you with a file or a content, analyze its content, and transform it into a visually appealing and well-structured webpage.### Content Requirements* Maintain the core information from the original file while presenting it in a clearer and more visually engaging format.⠀Design Style* Follow a modern and minimalistic design inspired by Linear App.* Use a clear visual hierarchy to emphasize important content.* Adopt a professional and harmonious color scheme that is easy on the eyes for extended reading.⠀Technical Specifications* Use HTML5, TailwindCSS 3.0+ (via CDN), and necessary JavaScript.* Implement a fully functional dark/light mode toggle, defaulting to the system setting.* Ensure clean, well-structured code with appropriate comments for easy understanding and maintenance.⠀Responsive Design* The page must be fully responsive, adapting seamlessly to mobile, tablet, and desktop screens.* Optimize layout and typography for different screen sizes.* Ensure a smooth and intuitive touch experience on mobile devices.⠀Icons & Visual Elements* Use professional icon libraries like Font Awesome or Material Icons (via CDN).* Integrate illustrations or charts that best represent the content.* Avoid using emojis as primary icons.* Check if any icons cannot be loaded.⠀User Interaction & ExperienceEnhance the user experience with subtle micro-interactions:* Buttons should have slight enlargement and color transitions on hover.* Cards should feature soft shadows and border effects on hover.* Implement smooth scrolling effects throughout the page.* Content blocks should have an elegant fade-in animation on load.⠀Performance Optimization* Ensure fast page loading by avoiding large, unnecessary resources.* Use modern image formats (WebP) with proper compression.* Implement lazy loading for content-heavy pages.⠀Output Requirements* Deliver a fully functional standalone HTML file, including all necessary CSS and JavaScript.* Ensure the code meets W3C standards with no errors or warnings.* Maintain consistent design and functionality across different browsers.⠀Create the most effective and visually appealing webpage based on the uploaded file's content type (document, data, images, etc.).Your output is only one HTML file, do not present any other notes on the HTML."
        
        # Estimated system prompt tokens (if exact count not available)
        system_prompt_tokens = len(system_prompt) // 3
        
        # Estimate tokens in content
        content_tokens = len(content) // 4
        
        # Total estimated tokens
        estimated_tokens = system_prompt_tokens + content_tokens
        
        # Calculate estimated cost (as of current pricing)
        estimated_cost = (estimated_tokens / 1000000) * 3.0  # $3 per million tokens
        
        # Add thinking budget in cost estimate
        thinking_budget = int(data.get('thinking_budget', DEFAULT_THINKING_BUDGET))
        if thinking_budget > 0:
            estimated_cost += (thinking_budget / 1000000) * 3.0  # Add thinking cost
        
        # Calculate max safe input tokens
        max_safe_input_tokens = 200000  # Claude 3.7 context window
        
        return jsonify({
            'estimated_tokens': estimated_tokens,
            'estimated_cost': round(estimated_cost, 6),
            'max_safe_input_tokens': max_safe_input_tokens
        })
    except Exception as e:
        return jsonify({"error": f"Error analyzing tokens: {str(e)}"}), 500

# Define helper functions for streaming
def format_stream_event(event_type, data=None):
    """Format a Server-Sent Event (SSE) message"""
    buffer = f"event: {event_type}\n"
    if data:
        # For status events, expose dispatch-friendly format
        if event_type == "status":
            buffer += f"data: {json.dumps(data)}\n"
            # Add a special field to dispatch custom event on the client side
            buffer += f"id: status_{int(time.time())}\n"
            buffer += f"retry: 15000\n"  # Tell client to retry connection after 15 seconds if dropped
        
        # For error events, add enough info for the client to handle it
        elif event_type == "error":
            # Make sure error data includes code if available
            if isinstance(data, dict) and not data.get("code") and "details" in data:
                # Try to extract code from details if it's a JSON string
                try:
                    details = data["details"]
                    if isinstance(details, str) and "{" in details and "code" in details:
                        import re
                        code_match = re.search(r'"code"\s*:\s*(\d+)', details)
                        if code_match:
                            data["code"] = int(code_match.group(1))
                except Exception:
                    pass  # Ignore any errors in code extraction
            
            buffer += f"data: {json.dumps(data)}\n"
            # Add a special field to dispatch custom event
            buffer += f"id: error_{int(time.time())}\n"
        else:
            # Regular event
            buffer += f"data: {json.dumps(data)}\n"
    buffer += "\n"
    return buffer

def create_stream_generator(client, system_prompt, user_message, model, max_tokens, temperature, thinking_budget=None):
    """Create a generator that yields SSE events for streaming Claude responses"""
    try:
        # Create message parameters
        message_params = {
            "model": model,
            "max_tokens": max_tokens,
            "temperature": temperature,
            "system": system_prompt,
            "messages": [
                {
                    "role": "user",
                    "content": user_message
                }
            ],
            "stream": True,
            "anthropic_metadata": {
                "user_session_id": f"file-visualizer-{int(time.time())}",
            }
        }
        
        # Add thinking parameter if specified
        if thinking_budget and thinking_budget > 0:
            try:
                # For newer versions of the Anthropic library
                from anthropic.types import Thinking
                message_params["thinking"] = Thinking(
                    enabled=True,
                    budget_limit=thinking_budget
                )
            except ImportError:
                # For older versions or fallback
                message_params["thinking"] = {
                    "enabled": True,
                    "budget_limit": thinking_budget
                }
        
        # Start the streaming response
        yield format_stream_event("stream_start", {"message": "Stream starting"})
        
        # Create streaming API call
        with client.messages.stream(**message_params) as stream:
            for chunk in stream:
                # Handle thinking updates
                if chunk.type == "thinking":
                    thinking_data = {
                        "type": "thinking_update",
                        "chunk_id": stream.message.id,
                        "thinking": {
                            "content": chunk.thinking.content if hasattr(chunk.thinking, "content") else ""
                        }
                    }
                    yield format_stream_event("content", thinking_data)
                
                # Handle content block deltas (the actual generated text)
                elif chunk.type == "content_block_delta":
                    content_data = {
                        "type": "content_block_delta",
                        "chunk_id": stream.message.id,
                        "delta": {
                            "text": chunk.delta.text
                        }
                    }
                    yield format_stream_event("content", content_data)
                
                # Handle message complete
                elif chunk.type == "message_complete":
                    usage_data = None
                    if hasattr(stream.message, "usage"):
                        usage_data = {
                            "input_tokens": stream.message.usage.input_tokens,
                            "output_tokens": stream.message.usage.output_tokens
                        }
                        
                        # Add thinking tokens if available
                        if hasattr(stream.message.usage, "thinking_tokens"):
                            usage_data["thinking_tokens"] = stream.message.usage.thinking_tokens
                    
                    complete_data = {
                        "type": "message_complete",
                        "message_id": stream.message.id,
                        "chunk_id": stream.message.id,
                        "usage": usage_data
                    }
                    yield format_stream_event("content", complete_data)
            
            # End of stream event
            yield format_stream_event("stream_end", {"message": "Stream complete"})
                
    except Exception as e:
        error_data = {
            "type": "error",
            "error": str(e)
        }
        yield format_stream_event("error", error_data)

# Add API endpoint for streaming processing
@app.route('/api/process-stream', methods=['POST'])
def process_stream():
    """
    Process a streaming request with reconnection support.
    """
    # Extract request data
    data = request.get_json()
    api_key = data.get('api_key')
    
    # Check for file upload fields
    file_name = data.get('file_name', '')
    file_content = data.get('file_content', '')
    
    # Check for both 'content' and 'source' parameters for compatibility
    content = data.get('content', '')
    if not content:
        content = data.get('source', '')  # Fallback to 'source' if 'content' is empty
    
    # Handle file content if provided
    if file_name and file_content:
        try:
            # Extract file extension
            file_ext = file_name.split('.')[-1].lower() if '.' in file_name else 'txt'
            print(f"Processing uploaded file: {file_name} with extension {file_ext}")
            
            # Create a temporary file
            temp_file_path = f"/tmp/{file_name}"
            
            # For binary files (PDF, DOCX, etc.), decode base64
            try:
                file_content_bytes = base64.b64decode(file_content)
                print(f"Successfully decoded base64 content, size: {len(file_content_bytes)} bytes")
            except Exception as decode_error:
                print(f"Error decoding base64 content: {str(decode_error)}")
                # Try to fix padding if that's the issue
                padded_content = file_content + '=' * (4 - len(file_content) % 4) if len(file_content) % 4 != 0 else file_content
                file_content_bytes = base64.b64decode(padded_content)
                print(f"Successfully decoded base64 content after padding fix, size: {len(file_content_bytes)} bytes")
            
            with open(temp_file_path, 'wb') as f:
                f.write(file_content_bytes)
            print(f"Wrote temporary file to {temp_file_path}")
            
            # Process the file based on type
            if file_ext == 'pdf':
                # Process PDF
                print(f"Processing PDF file")
                reader = PyPDF2.PdfReader(temp_file_path)
                text_content = ""
                for page_num in range(len(reader.pages)):
                    text_content += reader.pages[page_num].extract_text() + "\n"
                content = text_content
                print(f"Extracted {len(content)} characters from PDF")
                
            elif file_ext in ['docx', 'doc']:
                # Process Word document
                print(f"Processing Word document")
                doc = docx.Document(temp_file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                print(f"Extracted {len(content)} characters from Word document")
                
            else:
                # For text-based files, assume it's already decoded properly
                print(f"Processing text-based file")
                with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                print(f"Read {len(content)} characters from text file")
                
            print(f"Successfully processed uploaded file: {file_name}, extracted {len(content)} characters")
            
        except Exception as e:
            error_msg = f"Error processing file upload: {str(e)}"
            print(error_msg)
            traceback_str = traceback.format_exc()
            print(f"Traceback: {traceback_str}")
            return jsonify({"success": False, "error": error_msg}), 400
    
    # If both are empty, return an error
    if not content:
        return jsonify({"success": False, "error": "Source code or text is required"}), 400
    
    format_prompt = data.get('format_prompt', '')
    model = data.get('model', 'claude-3-7-sonnet-20250219')  # Updated to Claude 3.7
    max_tokens = int(data.get('max_tokens', DEFAULT_MAX_TOKENS))
    temperature = float(data.get('temperature', 0.5))
    thinking_budget = int(data.get('thinking_budget', DEFAULT_THINKING_BUDGET))
    
    # Reconnection support
    session_id = data.get('session_id', str(uuid.uuid4()))
    is_reconnect = data.get('is_reconnect', False)
    last_chunk_id = data.get('last_chunk_id', None)
    
    # Check if we have a cached response for this session
    if is_reconnect and session_id in session_cache:
        cached_data = session_cache[session_id]
        app.logger.info(f"Found cached data for session {session_id}, resuming from chunk {last_chunk_id}")
        
        # If we have partial content already generated, use that to save time
        if 'generated_text' in cached_data:
            # Use streaming response to deliver cached content and then continue
            return Response(
                stream_with_context(resume_from_cache(session_id, last_chunk_id, api_key)),
                content_type='text/event-stream'
            )
    
    # Create Anthropic client
    client = None
    try:
        client = create_anthropic_client(api_key)
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"API key validation failed: {str(e)}"
        })
    
    # Initialize session cache for this request
    session_cache[session_id] = {
        'created_at': time.time(),
        'last_updated': time.time(),
        'html_segments': [],
        'generated_text': '',
        'chunk_count': 0,
        'user_content': content[:100000],  # Store for potential reconnection
        'format_prompt': format_prompt,
        'model': model,
        'max_tokens': max_tokens,
        'temperature': temperature
    }
    
    # Prepare system prompt
    system_prompt = "I will provide you with a file or a content, analyze its content, and transform it into a visually appealing and well-structured webpage.### Content Requirements* Maintain the core information from the original file while presenting it in a clearer and more visually engaging format.⠀Design Style* Follow a modern and minimalistic design inspired by Linear App.* Use a clear visual hierarchy to emphasize important content.* Adopt a professional and harmonious color scheme that is easy on the eyes for extended reading.⠀Technical Specifications* Use HTML5, TailwindCSS 3.0+ (via CDN), and necessary JavaScript.* Implement a fully functional dark/light mode toggle, defaulting to the system setting.* Ensure clean, well-structured code with appropriate comments for easy understanding and maintenance.⠀Responsive Design* The page must be fully responsive, adapting seamlessly to mobile, tablet, and desktop screens.* Optimize layout and typography for different screen sizes.* Ensure a smooth and intuitive touch experience on mobile devices.⠀Icons & Visual Elements* Use professional icon libraries like Font Awesome or Material Icons (via CDN).* Integrate illustrations or charts that best represent the content.* Avoid using emojis as primary icons.* Check if any icons cannot be loaded.⠀User Interaction & ExperienceEnhance the user experience with subtle micro-interactions:* Buttons should have slight enlargement and color transitions on hover.* Cards should feature soft shadows and border effects on hover.* Implement smooth scrolling effects throughout the page.* Content blocks should have an elegant fade-in animation on load.⠀Performance Optimization* Ensure fast page loading by avoiding large, unnecessary resources.* Use modern image formats (WebP) with proper compression.* Implement lazy loading for content-heavy pages.* For large outputs, make sure the HTML can be incrementally rendered and uses efficient DOM structures.⠀Output Requirements* Deliver a fully functional standalone HTML file, including all necessary CSS and JavaScript.* Ensure the code meets W3C standards with no errors or warnings.* Maintain consistent design and functionality across different browsers.* Your output is only one HTML file, do not present any other notes on the HTML. Also, try your best to visualize the whole content.⠀Create the most effective and visually appealing webpage based on the uploaded file's content type (document, data, images, etc.)."
    
    # Enhanced system prompt for large content handling
    if len(content) > 50000:  # If content is large
        system_prompt += "\n\nIMPORTANT: This is a large document. To ensure the generated HTML can be efficiently processed and rendered by browsers, please follow these additional guidelines:\n1. Implement progressive rendering techniques\n2. Minimize deep DOM nesting - keep DOM depth under 20 levels\n3. Use document fragments and lazy loading where appropriate\n4. Break large content into smaller sections using pagination or tabs\n5. Break large tables into smaller sections with pagination\n6. Use efficient CSS selectors (avoid descendant selectors when possible)\n7. Minimize JavaScript interactions and DOM manipulations\n8. Avoid complex CSS animations and transitions\n9. Use lightweight, optimized SVG instead of heavy images\n10. Implement lazy-loaded images with low-resolution placeholders\n11. Break long sections of text into separate elements with reasonable length"

    # If the content is extremely large, add even more constraints
    if len(content) > 100000:
        system_prompt += "\nEXTREMELY LARGE CONTENT DETECTED: Break the content into multiple pages and implement a navigation system. Do not use complex or heavy JavaScript frameworks. Keep CSS minimal and efficient."
    
    # Prepare user prompt - limit content size to avoid timeouts
    content_limit = min(len(content), 100000)  # Limit to 100k characters
    user_content = f"""
    {format_prompt}
    
    Here is the content to transform into a website:
    
    {content[:content_limit]}
    """
    
    # Define a streaming response generator with specific Claude 3.7 implementation
    def stream_generator():
        try:
            yield format_stream_event("stream_start", {"message": "Stream starting", "session_id": session_id})
            
            # Add retry logic with exponential backoff
            max_retries = MAX_RETRIES
            retry_count = 0
            backoff_time = MIN_BACKOFF_DELAY  # Start with minimum delay
            
            while retry_count <= max_retries:
                try:
                    # Use the Claude 3.7 specific implementation with beta parameter
                    with client.beta.messages.stream(
                        model="claude-3-7-sonnet-20250219",
                        max_tokens=max_tokens,
                        temperature=temperature,
                        system=system_prompt,
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": user_content
                                    }
                                ]
                            }
                        ],
                        thinking={
                            "type": "enabled",
                            "budget_tokens": thinking_budget
                        },
                        betas=[OUTPUT_128K_BETA],  # Using betas parameter instead of headers
                    ) as stream:
                        message_id = str(uuid.uuid4())
                        generated_text = ""
                        start_time = time.time()
                        chunk_count = 0
                        
                        # Store accumulated HTML for segmented delivery
                        html_segments = []
                        current_segment = ""
                        current_segment_size = 0
                        segment_counter = 0
                        max_segment_size = MAX_SEGMENT_SIZE  # 16KB per segment
                        
                        # Add checkpoint tracking
                        last_checkpoint_time = time.time()
                        checkpoint_counter = 0
                        
                        for chunk in stream:
                            # Handle potential disconnection by saving state frequently
                            try:
                                current_time = time.time()
                                chunk_count += 1
                                
                                # Update session cache with current progress
                                if session_id in session_cache:
                                    session_cache[session_id]['last_updated'] = current_time
                                    session_cache[session_id]['chunk_count'] = chunk_count
                                
                                # More frequent keepalive messages (every STREAM_CHUNK_SIZE chunks)
                                if chunk_count % STREAM_CHUNK_SIZE == 0:
                                    yield format_stream_event("keepalive", {
                                        "timestamp": current_time,
                                        "session_id": session_id,
                                        "chunk_count": chunk_count
                                    })
                                
                                # Handle thinking updates
                                if hasattr(chunk, "thinking") and chunk.thinking:
                                    thinking_data = {
                                        "type": "thinking_update",
                                        "chunk_id": f"{message_id}_{chunk_count}",
                                        "thinking": {
                                            "content": chunk.thinking.content if hasattr(chunk.thinking, "content") else ""
                                        }
                                    }
                                    yield format_stream_event("content", thinking_data)
                                
                                # Handle content block deltas (the actual generated text)
                                if hasattr(chunk, "delta") and hasattr(chunk.delta, "text"):
                                    delta_text = chunk.delta.text
                                    generated_text += delta_text
                                    
                                    # Update session cache with generated text
                                    if session_id in session_cache:
                                        session_cache[session_id]['generated_text'] = generated_text
                                    
                                    # Check if we need to create a checkpoint (every 2 minutes)
                                    if current_time - last_checkpoint_time > CHECKPOINT_INTERVAL:
                                        checkpoint_id = f"cp_{session_id}_{checkpoint_counter}"
                                        checkpoint_counter += 1
                                        last_checkpoint_time = current_time
                                        
                                        # Store checkpoint in the session cache
                                        session_cache[session_id]["checkpoints"] = session_cache[session_id].get("checkpoints", {})
                                        session_cache[session_id]["checkpoints"][checkpoint_id] = {
                                            "html_so_far": generated_text,
                                            "chunk_id": f"{message_id}_{chunk_count}",
                                            "timestamp": current_time,
                                            "chunk_count": chunk_count
                                        }
                                        
                                        # Send a checkpoint event
                                        yield format_stream_event("status", {
                                            "type": "checkpoint",
                                            "checkpoint_id": checkpoint_id,
                                            "timestamp": current_time,
                                            "chunk_id": f"{message_id}_{chunk_count}",
                                            "chunk_count": chunk_count,
                                            "message": "Progress checkpoint created"
                                        })
                                    
                                    # Build up the current segment
                                    current_segment += delta_text
                                    current_segment_size += len(delta_text)
                                    
                                    # Check if we should close and send this segment
                                    # We send segments when they reach max size or contain complete HTML tags
                                    if (current_segment_size >= max_segment_size or 
                                        (current_segment_size > 256 and  # Reduced from 512 bytes to 256 bytes
                                         (delta_text.endswith('</div>') or 
                                          delta_text.endswith('</section>') or
                                          delta_text.endswith('</p>') or
                                          delta_text.endswith('</table>') or
                                          delta_text.endswith('</li>') or
                                          delta_text.endswith('</h1>') or
                                          delta_text.endswith('</h2>') or
                                          delta_text.endswith('</h3>') or
                                          delta_text.endswith('</html>')))):
                                        
                                        # Store this segment
                                        html_segments.append(current_segment)
                                        segment_counter += 1
                                        
                                        # Update session cache with segments
                                        if session_id in session_cache:
                                            session_cache[session_id]['html_segments'] = html_segments.copy()
                                        
                                        # Send segment to client
                                        content_data = {
                                            "type": "content_block_delta",
                                            "chunk_id": f"{message_id}_{chunk_count}",
                                            "delta": {
                                                "text": current_segment
                                            },
                                            "segment": segment_counter,
                                            "session_id": session_id,
                                            "chunk_count": chunk_count
                                        }
                                        yield format_stream_event("content", content_data)
                                        
                                        # Send a keepalive after every segment to maintain connection
                                        yield format_stream_event("keepalive", {
                                            "timestamp": time.time(),
                                            "session_id": session_id,
                                            "chunk_count": chunk_count,
                                            "segment": segment_counter
                                        })
                                        
                                        # Reset for next segment
                                        current_segment = ""
                                        current_segment_size = 0
                                        
                                        # Add a short sleep to let the browser process
                                        if segment_counter % 3 == 0:  # Reduced from 5 to 3
                                            time.sleep(0.05)
                                    
                                    # For smaller updates, send frequently to maintain connection
                                    # Send even small updates every 2 chunks (reduced from 5)
                                    elif chunk_count % 2 == 0 and current_segment:
                                        content_data = {
                                            "type": "content_block_delta",
                                            "chunk_id": f"{message_id}_{chunk_count}",
                                            "delta": {
                                                "text": current_segment
                                            },
                                            "partial": True,
                                            "session_id": session_id,
                                            "chunk_count": chunk_count
                                        }
                                        yield format_stream_event("content", content_data)
                                
                            except (ConnectionError, BrokenPipeError) as e:
                                app.logger.error(f"Client disconnected during streaming: {str(e)}")
                                # Save the current state for potential reconnection
                                app.logger.warning(f"Saving state at chunk {chunk_count} for session {session_id}")
                                
                                # Make sure session cache is updated before breaking
                                if session_id in session_cache:
                                    session_cache[session_id]['generated_text'] = generated_text
                                    session_cache[session_id]['html_segments'] = html_segments.copy()
                                    session_cache[session_id]['chunk_count'] = chunk_count
                                break
                        
                        # If we have any remaining segment, send it
                        if current_segment:
                            html_segments.append(current_segment)
                            segment_counter += 1
                            content_data = {
                                "type": "content_block_delta",
                                "chunk_id": f"{message_id}_{chunk_count}",
                                "delta": {
                                    "text": current_segment
                                },
                                "segment": segment_counter,
                                "session_id": session_id,
                                "chunk_count": chunk_count
                            }
                            yield format_stream_event("content", content_data)
                        
                        # If we completed the stream successfully and have content
                        if len(generated_text) > 0:
                            # Stream completed successfully, break out of retry loop
                            break
                        else:
                            # If we broke out of the loop due to connection issue but have partial results
                            # Log the state for reconnection
                            app.logger.warning(f"Partial completion for session {session_id}, chunk count: {chunk_count}")
                            # Don't break here, let it retry if needed
                
                except Exception as e:
                    error_str = str(e)
                    error_details = ""
                    
                    # Check if it's an API error with a response
                    if hasattr(e, 'response') and hasattr(e.response, 'json'):
                        try:
                            error_details = e.response.json()
                            app.logger.error(f"API Error details: {error_details}")
                            
                            # Check specifically for overloaded error (529)
                            if isinstance(error_details, dict) and error_details.get('code') == 529:
                                if retry_count < max_retries:
                                    retry_count += 1
                                    
                                    # Calculate backoff with jitter to prevent thundering herd
                                    jitter = random.uniform(0.8, 1.2)
                                    wait_time = min(backoff_time * jitter, MAX_BACKOFF_DELAY)
                                    backoff_time = min(backoff_time * BACKOFF_FACTOR, MAX_BACKOFF_DELAY)
                                    
                                    app.logger.warning(f"Anthropic API overloaded. Retry {retry_count}/{max_retries} after {wait_time:.2f}s")
                                    yield format_stream_event("status", {
                                        "type": "status", 
                                        "message": f"Anthropic API temporarily overloaded. Retrying in {wait_time:.1f}s (attempt {retry_count}/{max_retries})...",
                                        "session_id": session_id,
                                        "retry": retry_count,
                                        "max_retries": max_retries
                                    })
                                    
                                    time.sleep(wait_time)
                                    continue  # Try again
                                else:
                                    app.logger.error(f"Max retries ({max_retries}) exceeded for API overload")
                                    yield format_stream_event("error", {
                                        "type": "error",
                                        "error": "Maximum retry attempts exceeded. Please try again later.",
                                        "details": "The AI service is currently experiencing high load. Your request could not be completed after multiple attempts.",
                                        "code": 529,
                                        "session_id": session_id
                                    })
                                    return
                        except Exception as json_err:
                            app.logger.error(f"Failed to parse error response: {str(json_err)}")
                    
                    # For other errors that are not 529
                    app.logger.error(f"Error in stream_generator: {error_str}")
                    if error_details:
                        app.logger.error(f"Error details: {error_details}")
                    
                    # Yield error and exit
                    yield format_stream_event("error", {
                        "type": "error",
                        "error": error_str,
                        "details": str(error_details),
                        "session_id": session_id
                    })
                    return
            
            # Send message complete event with usage statistics when available
            try:
                usage_data = None
                if hasattr(stream, "usage"):
                    usage_data = {
                        "input_tokens": stream.usage.input_tokens if hasattr(stream.usage, "input_tokens") else 0,
                        "output_tokens": stream.usage.output_tokens if hasattr(stream.usage, "output_tokens") else 0
                    }
                    # Calculate cost according to Anthropic pricing
                    usage_data["total_cost"] = (usage_data["input_tokens"] / 1000000 * 3.0) + (usage_data["output_tokens"] / 1000000 * 15.0)
                else:
                    # If usage is not available from stream, calculate manually
                    system_prompt_tokens = len(system_prompt) // 3
                    content_tokens = len(user_content) // 4
                    output_tokens = len(generated_text) // 4
                    
                    usage_data = {
                        "input_tokens": system_prompt_tokens + content_tokens,
                        "output_tokens": output_tokens,
                        "time_elapsed": round(time.time() - start_time, 2),
                        "total_cost": (system_prompt_tokens + content_tokens) / 1000000 * 3.0 + output_tokens / 1000000 * 15.0
                    }
                
                # Mark this session as complete in the cache
                if session_id in session_cache:
                    session_cache[session_id]['complete'] = True
                    session_cache[session_id]['usage'] = usage_data
                
                complete_data = {
                    "type": "message_complete",
                    "message_id": message_id,
                    "chunk_id": f"{message_id}_{chunk_count}",
                    "usage": usage_data,
                    "html": generated_text,
                    "session_id": session_id,
                    "final_chunk_count": chunk_count,
                    "segment_count": segment_counter
                }
                yield format_stream_event("content", complete_data)
                yield format_stream_event("stream_end", {"message": "Stream complete", "session_id": session_id})
            except (ConnectionError, BrokenPipeError) as e:
                app.logger.error(f"Client disconnected during completion: {str(e)}")
        except Exception as e:
            app.logger.error(f"Unexpected error in stream generator: {str(e)}")
            # Include stack trace for better debugging
            app.logger.error(traceback.format_exc())
            yield format_stream_event("error", {
                "type": "error",
                "error": str(e),
                "details": traceback.format_exc(),
                "session_id": session_id
            })
    
    # Generator for resuming from cache
    def resume_from_cache(session_id, last_chunk_id, api_key):
        try:
            # Get cached data
            cached_data = session_cache[session_id]
            app.logger.info(f"Resuming from cache for session {session_id}")
            
            # Start with a stream_start event
            yield format_stream_event("stream_start", {
                "message": "Resuming stream",
                "session_id": session_id,
                "is_resumed": True
            })
            
            # Extract segment number from last_chunk_id
            last_segment = 0
            if last_chunk_id and '_' in last_chunk_id:
                try:
                    last_segment = int(last_chunk_id.split('_')[1]) // 10  # Approximate segment
                except ValueError:
                    pass
            
            # Send all cached segments after the last known segment
            html_segments = cached_data.get('html_segments', [])
            
            for i, segment in enumerate(html_segments):
                segment_num = i + 1
                if segment_num > last_segment:
                    content_data = {
                        "type": "content_block_delta",
                        "chunk_id": f"{session_id}_{segment_num * 10}",
                        "delta": {
                            "text": segment
                        },
                        "segment": segment_num,
                        "session_id": session_id,
                        "chunk_count": segment_num * 10,
                        "is_cached": True
                    }
                    yield format_stream_event("content", content_data)
                    time.sleep(0.1)  # Short delay between segments
            
            # If generation was complete, send the completion event
            if cached_data.get('complete', False):
                app.logger.info(f"Sending cached completion for session {session_id}")
                
                complete_data = {
                    "type": "message_complete",
                    "message_id": session_id,
                    "chunk_id": f"{session_id}_{len(html_segments) * 10}",
                    "usage": cached_data.get('usage', {}),
                    "html": cached_data.get('generated_text', ''),
                    "session_id": session_id,
                    "final_chunk_count": len(html_segments) * 10,
                    "segment_count": len(html_segments),
                    "is_cached": True
                }
                yield format_stream_event("content", complete_data)
                yield format_stream_event("stream_end", {
                    "message": "Stream complete (from cache)",
                    "session_id": session_id
                })
                return
            
            # If generation was not complete, continue with the regular generator
            # Create a client and continue where we left off
            client = create_anthropic_client(api_key)
            
            # Continue with a new request
            app.logger.info(f"Continuing generation for session {session_id}")
            
            # Continue by creating a new generator
            content = cached_data.get('user_content', '')
            format_prompt = cached_data.get('format_prompt', '')
            max_tokens = cached_data.get('max_tokens', DEFAULT_MAX_TOKENS)
            temperature = cached_data.get('temperature', 0.5)
            
            # Add a message that we're continuing generation
            yield format_stream_event("status", {
                "type": "status",
                "message": "Continuing generation...",
                "session_id": session_id
            })
            
            # Continue with new generator - simplified to avoid nesting too deep
            # In production, this should be refactored to avoid code duplication
            yield from stream_generator()
            
        except Exception as e:
            app.logger.error(f"Error resuming from cache: {str(e)}")
            yield format_stream_event("error", {
                "type": "error",
                "error": f"Failed to resume: {str(e)}",
                "session_id": session_id
            })
            
            # Start a new generator from scratch as fallback
            yield format_stream_event("status", {
                "type": "status",
                "message": "Failed to resume. Starting new generation.",
                "session_id": session_id
            })
            
            # Re-initialize session
            session_cache[session_id] = {
                'created_at': time.time(),
                'last_updated': time.time(),
                'html_segments': [],
                'generated_text': '',
                'chunk_count': 0
            }
            
            yield from stream_generator()
    
    # Return streaming response
    response = Response(stream_with_context(stream_generator()), 
                         content_type='text/event-stream')
    response.headers['X-Accel-Buffering'] = 'no'  # Disable nginx buffering
    response.headers['Cache-Control'] = 'no-cache, no-transform'
    response.headers['Connection'] = 'keep-alive'
    response.headers['Keep-Alive'] = 'timeout=3600, max=2000'  # 60 minutes timeout (increased from 30)
    response.headers['X-Accel-Limit-Rate'] = '0'  # Disable rate limiting
    return response

# Clean up expired sessions from cache
@app.before_request
def cleanup_session_cache():
    current_time = time.time()
    expired_sessions = []
    
    for session_id, session_data in session_cache.items():
        # Clean up sessions older than SESSION_CACHE_EXPIRY
        if current_time - session_data.get('created_at', 0) > SESSION_CACHE_EXPIRY:
            expired_sessions.append(session_id)
    
    # Remove expired sessions
    for session_id in expired_sessions:
        del session_cache[session_id]
        
    # Log cleanup if sessions were removed
    if expired_sessions:
        app.logger.info(f"Cleaned up {len(expired_sessions)} expired sessions from cache")

# Add a simple test endpoint
@app.route('/api/test', methods=['GET', 'POST'])
def test_api():
    """Simple test endpoint to verify API connectivity"""
    if request.method == 'POST':
        data = request.get_json() or {}
        return jsonify({
            "status": "success", 
            "message": "API test endpoint is working",
            "received_data": data,
            "timestamp": time.time()
        })
    else:
        return jsonify({
            "status": "success", 
            "message": "API test endpoint is working", 
            "timestamp": time.time()
        })

@app.route('/api/test-generate', methods=['POST'])
def test_generate():
    """
    Test endpoint that uses minimal Anthropic API tokens to test functionality
    """
    try:
        # Get data from request
        data = request.get_json()
        content = data.get('content', '') or data.get('source', '')
        api_key = data.get('api_key')
        
        if not api_key:
            return jsonify({"success": False, "error": "API key is required"}), 400
            
        if not content:
            return jsonify({"success": False, "error": "Content is required"}), 400
        
        # Truncate content to just 500 characters to minimize token usage
        truncated_content = content[:500] + "..." if len(content) > 500 else content
        
        # Create Anthropic client
        client = None
        try:
            client = create_anthropic_client(api_key)
            print("Client created successfully")
        except Exception as e:
            print(f"Error creating client: {str(e)}")
            return jsonify({
                "success": False,
                "error": f"API key validation failed: {str(e)}"
            }), 400
        
        # Brief system prompt to minimize token usage
        system_prompt = "Generate minimal HTML to confirm the API integration works."
        
        # Minimal user prompt
        user_content = f"Generate a very simple HTML page that confirms the API works. The content is: {truncated_content}"
        
        # Record start time
        start_time = time.time()
        
        # Implement retry logic with exponential backoff
        max_retries = 5
        retry_count = 0
        base_delay = 1  # Start with 1 second (in seconds)
        success = False
        response = None
        html_output = None
        error_message = None
        
        while retry_count <= max_retries and not success:
            try:
                print(f"Calling Anthropic API with test message (attempt {retry_count + 1}/{max_retries + 1})")
                # Call Anthropic API with minimal settings
                response = client.messages.create(
                    model="claude-3-haiku-20240307",  # Use smaller model to save tokens
                    max_tokens=100,  # Minimal output
                    temperature=0.5,
                    system=system_prompt,
                    messages=[
                        {
                            "role": "user",
                            "content": user_content
                        }
                    ]
                )
                
                # If we get here without error, set success flag
                success = True
                
                # Get the response content if successful
                if hasattr(response, 'content') and len(response.content) > 0:
                    if hasattr(response.content[0], 'text'):
                        html_output = response.content[0].text
                    elif isinstance(response.content[0], dict) and 'text' in response.content[0]:
                        html_output = response.content[0]['text']
                
            except Exception as e:
                error_message = str(e)
                print(f"Error in API call attempt {retry_count + 1}: {error_message}")
                
                # Check if this is an overloaded error (529)
                is_overloaded = "529" in error_message and "Overloaded" in error_message
                
                if is_overloaded and retry_count < max_retries:
                    retry_count += 1
                    wait_time = base_delay * (2 ** (retry_count - 1))  # Exponential backoff
                    print(f"Anthropic API overloaded. Retry {retry_count}/{max_retries} after {wait_time}s")
                    time.sleep(wait_time)
                    continue
                else:
                    # Break the loop for other errors or if retries exhausted
                    break
        
        # Calculate elapsed time
        elapsed_time = time.time() - start_time
        
        # If successful, return the generated HTML
        if success and html_output:
            # Calculate token usage (estimates)
            system_prompt_tokens = len(system_prompt) // 3
            content_tokens = len(truncated_content) // 4
            output_tokens = len(html_output) // 4
            
            # Return response with HTML and token usage statistics
            return jsonify({
                "success": True,
                "html": html_output,
                "usage": {
                    "input_tokens": system_prompt_tokens + content_tokens,
                    "output_tokens": output_tokens,
                    "time_elapsed": elapsed_time,
                    "total_cost": (system_prompt_tokens + content_tokens) / 1000000 * 3.0 + output_tokens / 1000000 * 15.0
                },
                "test_mode": True,
                "message": "Test completed with minimal token usage"
            })
        else:
            # Return an error message with nice HTML
            is_overloaded = error_message and "529" in error_message and "Overloaded" in error_message
            
            error_html = f"""
            <html>
            <head>
                <title>Test Mode - {("API Overloaded" if is_overloaded else "API Error")}</title>
                <style>
                    body {{ font-family: system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 2rem; max-width: 800px; margin: 0 auto; }}
                    .error {{ background-color: #fee2e2; border-left: 4px solid #ef4444; padding: 1rem; border-radius: 4px; margin: 1rem 0; }}
                    pre {{ background-color: #f7fafc; padding: 1rem; border-radius: 4px; overflow-x: auto; }}
                    h1 {{ color: #1e293b; }}
                </style>
            </head>
            <body>
                <h1>Test Mode - {("API Overloaded" if is_overloaded else "API Error")}</h1>
                <p>{"The Anthropic API is currently experiencing high demand and is overloaded. We tried several times but could not get a response." if is_overloaded else "There was an error calling the Anthropic API:"}</p>
                <div class="error">
                    <pre>{error_message or "Unknown error"}</pre>
                </div>
                <p>Please try again later or check your API key.</p>
            </body>
            </html>
            """
            
            return jsonify({
                "success": False,
                "error": error_message or "Error calling Anthropic API",
                "html": error_html,
                "code": 529 if is_overloaded else 500,
                "usage": {
                    "input_tokens": 0,
                    "output_tokens": 0,
                    "time_elapsed": elapsed_time,
                    "total_cost": 0
                },
                "retries_attempted": retry_count,
                "test_mode": True
            }), 200  # Return 200 even with error since we're providing valid HTML
            
    except Exception as e:
        print(f"Unhandled exception in test_generate: {str(e)}")
        error_html = f"""
        <html>
        <head>
            <title>Test Mode - Server Error</title>
            <style>
                body {{ font-family: system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 2rem; max-width: 800px; margin: 0 auto; }}
                .error {{ background-color: #fee2e2; border-left: 4px solid #ef4444; padding: 1rem; border-radius: 4px; }}
                h1 {{ color: #1e293b; }}
            </style>
        </head>
        <body>
            <h1>Test Mode - Server Error</h1>
            <p>An unexpected error occurred while processing your request:</p>
            <div class="error">
                <p>{str(e)}</p>
            </div>
            <p>Please try again later.</p>
        </body>
        </html>
        """
        
        return jsonify({
            "success": False, 
            "error": f"Error in test generation: {str(e)}",
            "html": error_html,
            "test_mode": True
        }), 200  # Return 200 for better client handling

# Add a new route for Gemini API processing
@app.route('/api/process-gemini', methods=['POST'])
def process_gemini():
    """
    Process a file using the Google Gemini API and return HTML.
    """
    # Get the data from the request
    print("\n==== API PROCESS GEMINI REQUEST RECEIVED ====")
    data = request.get_json()
    
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    # Extract the API key and content
    api_key = data.get('api_key')
    content = data.get('content')
    format_prompt = data.get('format_prompt', '')
    max_tokens = int(data.get('max_tokens', GEMINI_MAX_OUTPUT_TOKENS))
    temperature = float(data.get('temperature', GEMINI_TEMPERATURE))
    
    print(f"Processing Gemini request with max_tokens={max_tokens}, content_length={len(content) if content else 0}")
    
    # Check if we have the required data
    if not api_key or not content:
        return jsonify({'error': 'API key and content are required'}), 400
    
    # Check if Gemini is available
    if not GEMINI_AVAILABLE:
        return jsonify({
            'error': 'Google Generative AI package is not installed on the server.'
        }), 500
    
    try:
        # Use our helper function to create a Gemini client
        client = create_gemini_client(api_key)
        
        # Prepare user message with content and additional prompt
        user_content = content
        if format_prompt:
            user_content = f"{user_content}\n\n{format_prompt}"
        
        print("Creating Gemini model...")
        
        # Get the model
        model = client.get_model(GEMINI_MODEL)
        
        # Configure generation parameters with more reliable settings
        generation_config = {
            "max_output_tokens": max_tokens,
            "temperature": temperature,
            "top_p": GEMINI_TOP_P,
            "top_k": GEMINI_TOP_K
        }
        
        # Use more reliable safety settings to prevent empty responses
        safety_settings = {
            "harassment": "block_none",
            "hate_speech": "block_none",
            "sexual": "block_none",
            "dangerous": "block_none",
        }
        
        # Create the prompt
        prompt = f"""
{SYSTEM_INSTRUCTION}

Here is the content to transform into a website:

{user_content}
"""
        
        # Generate content
        print(f"Generating content with {GEMINI_MODEL}, max_tokens={max_tokens}, temperature={temperature}")
        
        # Use more robust error handling and retry logic
        max_retries = 3
        retry_count = 0
        last_error = None
        
        while retry_count < max_retries:
            try:
                response = model.generate_content(
                    prompt,
                    generation_config=generation_config,
                    safety_settings=safety_settings
                )
                
                # We got a response, break out of retry loop
                break
                
            except Exception as e:
                last_error = e
                retry_count += 1
                print(f"Gemini API error (attempt {retry_count}/{max_retries}): {str(e)}")
                
                # Wait before retrying
                if retry_count < max_retries:
                    wait_time = 2 ** retry_count  # Exponential backoff
                    print(f"Waiting {wait_time} seconds before retry...")
                    time.sleep(wait_time)
                else:
                    # Max retries reached, raise the last error
                    raise
        
        # Extract the HTML from the response
        html_content = ""
        
        # Try multiple approaches to extract content
        if hasattr(response, 'text'):
            html_content = response.text
        elif hasattr(response, 'parts') and response.parts:
            html_content = response.parts[0].text
        
        # Clean HTML content if it contains markdown-style code blocks
        if html_content and ('```html' in html_content or '```' in html_content):
            # Extract the actual HTML from between the markdown code blocks
            print("Detected markdown code blocks in Gemini response, extracting HTML...")
            
            # First try with ```html specific tag
            html_match = re.search(r'```html\s*([\s\S]*?)\s*```', html_content)
            if html_match and html_match[1]:
                html_content = html_match[1].strip()
                print(f"Extracted HTML from markdown code blocks, new length: {len(html_content)}")
            else:
                # Try with just ``` blocks
                html_match = re.search(r'```\s*([\s\S]*?)\s*```', html_content)
                if html_match and html_match[1]:
                    html_content = html_match[1].strip()
                    print(f"Extracted HTML from generic markdown blocks, new length: {len(html_content)}")
        
        # If we still don't have content, use string representation
        if not html_content:
            print("Warning: Could not extract HTML content using standard methods")
            html_content = str(response)
        
        # Verify that the content looks like HTML
        if not html_content.strip().startswith('<') and not '<html' in html_content:
            print(f"Warning: Content doesn't appear to be HTML. Content starts with: {html_content[:100]}")
            
            # Attempt to fix by wrapping in HTML tags if it's just text content
            if not html_content.strip().startswith('<'):
                print("Attempting to fix by wrapping in HTML tags")
                html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Generated Content</title>
    <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
</head>
<body class="bg-gray-100 dark:bg-gray-900 text-gray-800 dark:text-gray-200">
    <div class="container mx-auto p-4">
        {html_content}
    </div>
    <script>
        // Simple dark mode toggle
        if (window.matchMedia && window.matchMedia('(prefers-color-scheme: dark)').matches) {{
            document.documentElement.classList.add('dark');
        }}
    </script>
</body>
</html>
"""
        
        # Get usage stats (approximate since Gemini doesn't provide exact token counts)
        input_tokens = len(prompt.split()) * 1.3  # Rough estimate: ~1.3 tokens per word
        output_tokens = len(html_content.split()) * 1.3
        
        # Ensure we don't have zero tokens (minimum of source length / 3)
        input_tokens = max(len(content.split()) * 1.3, input_tokens)
        
        # Round to integers
        input_tokens = max(1, int(input_tokens))
        output_tokens = max(1, int(output_tokens))
        
        # Log response
        print(f"Successfully generated HTML with Gemini. Input tokens: {input_tokens}, Output tokens: {output_tokens}")
        
        # Return the response
        return jsonify({
            'html': html_content,
            'model': GEMINI_MODEL,
            'usage': {
                'input_tokens': input_tokens,
                'output_tokens': output_tokens,
                'total_tokens': input_tokens + output_tokens,
                'total_cost': 0.0  # Gemini API is currently free
            }
        })
    
    except Exception as e:
        error_message = str(e)
        traceback_str = traceback.format_exc()
        print(f"Error in /api/process-gemini: {error_message}")
        print(f"Traceback: {traceback_str}")
        
        # Create a graceful fallback error page as HTML
        error_html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Error Processing Content</title>
            <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
        </head>
        <body class="bg-gray-100 dark:bg-gray-900 text-gray-800 dark:text-gray-200">
            <div class="container mx-auto p-8 max-w-3xl">
                <div class="bg-white dark:bg-gray-800 p-6 rounded-lg shadow-lg">
                    <h1 class="text-2xl font-bold text-red-600 mb-4">Error Processing Content</h1>
                    <p class="mb-4">There was an error processing your content with the Gemini API:</p>
                    <div class="bg-gray-100 dark:bg-gray-700 p-4 rounded overflow-auto mb-4">
                        <code class="text-sm">{error_message}</code>
                    </div>
                    <p class="mb-2">Possible solutions:</p>
                    <ul class="list-disc pl-5 mb-4">
                        <li>Try again with a smaller content size</li>
                        <li>Check your Gemini API key</li>
                        <li>Try with a different content format</li>
                        <li>Switch to the Claude API if available</li>
                    </ul>
                </div>
            </div>
        </body>
        </html>
        """
        
        # Return error as both JSON and HTML
        return jsonify({
            'error': f'Server error: {error_message}', 
            'html': error_html
        }), 500

# Add a streaming endpoint for Gemini
@app.route('/api/process-gemini-stream', methods=['POST'])
def process_gemini_stream():
    """
    Process a streaming request using Google Gemini API.
    """
    # Extract request data
    data = request.get_json()
    api_key = data.get('api_key')
    
    # Check for both 'content' and 'source' parameters for compatibility
    content = data.get('content', '')
    if not content:
        content = data.get('source', '')  # Fallback to 'source' if 'content' is empty
    
    # If content is empty, return an error
    if not content:
        return jsonify({"success": False, "error": "Source code or text is required"}), 400
    
    format_prompt = data.get('format_prompt', '')
    max_tokens = int(data.get('max_tokens', GEMINI_MAX_OUTPUT_TOKENS))
    temperature = float(data.get('temperature', GEMINI_TEMPERATURE))
    
    # Reconnection support
    session_id = data.get('session_id', str(uuid.uuid4()))
    
    # Check if Gemini is available
    if not GEMINI_AVAILABLE:
        return jsonify({
            'error': 'Google Generative AI package is not installed on the server.'
        }), 500
    
    # Create Gemini client
    try:
        client = create_gemini_client(api_key)
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"API key validation failed: {str(e)}"
        })
    
    # Initialize session cache for this request
    session_cache[session_id] = {
        'created_at': time.time(),
        'last_updated': time.time(),
        'html_segments': [],
        'generated_text': '',
        'chunk_count': 0,
        'user_content': content[:100000],  # Store for potential reconnection
        'format_prompt': format_prompt,
        'model': GEMINI_MODEL,
        'max_tokens': max_tokens,
        'temperature': temperature
    }
    
    # Prepare prompt
    prompt = f"""
{SYSTEM_INSTRUCTION}

Here is the content to transform into a website:

{content[:100000]}
"""
    
    if format_prompt:
        prompt += f"\n\n{format_prompt}"
    
    # Define the streaming response generator
    def gemini_stream_generator():
        try:
            yield format_stream_event("stream_start", {"message": "Stream starting", "session_id": session_id})
            
            # Get the model
            model = client.get_model(GEMINI_MODEL)
            
            # Configure generation parameters
            generation_config = {
                "max_output_tokens": max_tokens,
                "temperature": temperature,
                "top_p": GEMINI_TOP_P,
                "top_k": GEMINI_TOP_K
            }
            
            # Generate content with streaming
            try:
                print(f"Starting Gemini content generation with model {GEMINI_MODEL}")
                print(f"Generation config: max_tokens={generation_config['max_output_tokens']}, temp={generation_config['temperature']}")
                
                # Use more reliable safety settings to prevent empty responses
                safety_settings = {
                    "harassment": "block_none",
                    "hate_speech": "block_none",
                    "sexual": "block_none",
                    "dangerous": "block_none",
                }
                
                # Generate content with the specified configurations
                stream_response = model.generate_content(
                    prompt,
                    generation_config=generation_config,
                    safety_settings=safety_settings,
                    stream=True
                )
                
                # Log success
                print("Successfully created Gemini stream response object")
                
                # Use our custom streaming response class
                with GeminiStreamingResponse(stream_response, session_id) as gemini_stream:
                    print(f"Entering GeminiStreamingResponse context with session ID: {session_id}")
                    chunk_count = 0
                    for chunk in gemini_stream:
                        chunk_count += 1
                        if chunk_count % 10 == 0:
                            print(f"Processed {chunk_count} chunks from Gemini stream")
                        yield chunk
                    
                    print(f"Completed streaming {chunk_count} chunks from Gemini")
                    # Stream end event
                    yield format_stream_event("stream_end", {
                        "message": "Stream complete",
                        "session_id": session_id
                    })
            except Exception as e:
                error_message = str(e)
                print(f"Error during Gemini content generation: {error_message}")
                traceback_str = traceback.format_exc()
                print(f"Traceback: {traceback_str}")
                
                yield format_stream_event("error", {
                    "type": "error",
                    "error": error_message,
                    "details": traceback_str,
                    "session_id": session_id
                })
            
        except Exception as e:
            error_message = str(e)
            print(f"Error in gemini_stream_generator: {error_message}")
            
            yield format_stream_event("error", {
                "type": "error",
                "error": error_message,
                "session_id": session_id
            })
    
    # Return the streaming response
    return Response(
        stream_with_context(gemini_stream_generator()),
        content_type='text/event-stream'
    )

@app.route('/api/version', methods=['GET'])
def get_version():
    """Return the application version information."""
    return jsonify({
        'version': __version__,
        'name': 'File Visualization',
        'gemini_available': GEMINI_AVAILABLE,
        'gemini_model': GEMINI_MODEL,
        'claude_model': 'claude-3-7-sonnet-20250219',
        'timestamp': datetime.now().isoformat()
    })

if __name__ == "__main__":
    # Parse command line arguments
    parser = argparse.ArgumentParser(description="Start the File Visualizer server")
    
    # Add command line arguments
    parser.add_argument("--port", type=int, default=5009, help="Port to run the server on (default: 5009)")
    parser.add_argument("--host", type=str, default="0.0.0.0", help="Host to run the server on (default: 0.0.0.0)")
    parser.add_argument("--debug", action="store_true", help="Run in debug mode")
    parser.add_argument("--no-reload", action="store_true", help="Disable auto-reload")
    
    # Parse arguments
    args = parser.parse_args()
    
    # Configure logging for better debugging
    logging.basicConfig(
        level=logging.INFO if not args.debug else logging.DEBUG,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Configure longer timeouts to handle large content
    from werkzeug.serving import run_simple
    
    # Extended timeout settings
    app.config['TIMEOUT'] = 900  # 15 minutes timeout for requests 
    
    # Find an available port if the specified one is in use
    port = args.port
    while True:
        try:
            # Check if the port is available
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.bind((args.host, port))
                break
        except OSError:
            # Port is in use, try the next one
            print(f"Port {port} is in use, trying {port + 1}...")
            port += 1
    
    print(f"\n==== File Visualizer running at http://{args.host if args.host != '0.0.0.0' else 'localhost'}:{port} ====\n")
    print(f"Visit http://localhost:{port} in your browser to use the application.")
    
    # Run with enhanced settings for larger content
    run_simple(
        args.host, 
        port, 
        app, 
        use_reloader=not args.no_reload,
        use_debugger=args.debug,
        threaded=True,
        passthrough_errors=args.debug
    )

# Important: Export the Flask app for Vercel
app